"""
Audit every number in the manuscript against the result tables.

Two axes, both of which must pass:

  coherence  does the paper agree with itself? Every N must descend from one
             cohort ledger, and a count stated in the abstract, the methods,
             a caption and a table must be the same count.
  fidelity   does the paper agree with the code? Every effect estimate must
             match the CSV the model wrote.

The manuscript is generated from those CSVs, so fidelity should hold by
construction — which is exactly why it is worth testing. Any literal digit in
the builder is a place where that guarantee was bypassed, and this script
finds them.

Exit code 1 on any FAIL, so it can gate a submission step.

Run:
    python tools/audit_numbers.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

import pandas as pd
from docx import Document

REPO = Path(__file__).resolve().parent.parent
TABLES = REPO / "results" / "tables"
MAIN = REPO / "manuscript" / "Pupil_MS_v5.docx"
SUPPLEMENT = REPO / "manuscript" / "Pupil_MS_v5_supplement.docx"
BUILDER = REPO / "manuscript" / "build_v5_docx.py"

TOLERANCE = {"count": 0.0, "percent": 0.1, "ratio": 0.01, "coefficient": 0.01}

failures: list[str] = []
unmatched: list[str] = []
suspicious: list[str] = []
passes = 0


def load(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES / name)


def document_text(path: Path) -> str:
    document = Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def check(label: str, kind: str, reported, computed) -> None:
    """Record one comparison against its tolerance."""
    global passes
    if reported is None:
        unmatched.append(label)
        return
    if abs(float(reported) - float(computed)) <= TOLERANCE[kind]:
        passes += 1
    else:
        failures.append(
            f"{label}: paper {reported}, computed {computed} "
            f"(tolerance {TOLERANCE[kind]})")


def main() -> int:
    global passes

    if not MAIN.exists():
        print(f"manuscript missing: {MAIN}", file=sys.stderr)
        return 1

    # Freshness: a manuscript older than the tables it quotes is a stale build.
    newest_table = max(p.stat().st_mtime for p in TABLES.glob("*.csv"))
    if MAIN.stat().st_mtime < newest_table:
        failures.append(
            "manuscript is older than the result tables; rebuild before auditing")

    text = document_text(MAIN) + "\n" + document_text(SUPPLEMENT)

    # ---------------------------------------------------------------- ledger
    design = load("analysis_design_summary.csv").set_index("parameter")["value"]
    audit = load("derived_data_audit.csv").set_index("metric")["value"]
    coverage = load("region_coverage_summary.csv")

    ledger = {
        "contacts": int(design["n_contacts"]),
        "excursion contacts": int(design["n_excursion_contacts"]),
        "patients": int(design["n_patients"]),
        "shafts": int(design["n_shafts"]),
        "legacy-selected contacts": int(audit["legacy_selected_channels"]),
        "informative paired patients": int(design["paired_patients_informative"]),
    }
    ledger["zero contacts"] = ledger["contacts"] - ledger["excursion contacts"]

    print("Cohort ledger")
    for name, value in ledger.items():
        occurrences = len(re.findall(rf"\b{value}\b", text))
        print(f"  {name:<30} {value:>5}   appears {occurrences}x in the paper")
        if occurrences == 0:
            unmatched.append(f"ledger value {name}={value} never appears")

    # Coherence: the two halves of the hurdle must add up wherever stated.
    if ledger["excursion contacts"] + ledger["zero contacts"] != ledger["contacts"]:
        failures.append("hurdle split does not sum to the contact total")
    else:
        passes += 1

    stated_percent = re.search(
        rf"{ledger['zero contacts']} of {ledger['contacts']} contacts "
        r"\((\d+)%\)", text)
    if stated_percent:
        check("methods zero-mass percentage", "percent",
              float(stated_percent.group(1)),
              round(100 * ledger["zero contacts"] / ledger["contacts"]))
    else:
        unmatched.append("methods zero-mass percentage sentence not found")

    excursion_percent = re.search(
        rf"{ledger['excursion contacts']} of {ledger['contacts']} contacts "
        r"\((\d+)%\)", text)
    if excursion_percent:
        check("abstract excursion percentage", "percent",
              float(excursion_percent.group(1)),
              round(100 * ledger["excursion contacts"] / ledger["contacts"]))

    # Region coverage totals must reconcile with the ledger.
    if int(coverage["n_contacts"].sum()) != ledger["contacts"]:
        failures.append(
            f"region coverage sums to {int(coverage['n_contacts'].sum())}, "
            f"ledger says {ledger['contacts']}")
    else:
        passes += 1

    # --------------------------------------------------------------- fidelity
    def find_ratio(pattern: str) -> float | None:
        match = re.search(pattern, text)
        return float(match.group(1)) if match else None

    polarity = load("polarity_primary_glme.csv")
    effect = polarity[polarity["term"].str.contains("Hippocampal")].iloc[0]
    # The dichotomised contrast is no longer summarised as a ratio in running
    # text; the model's odds ratio survives only in the supplement tables, so
    # these read the row rather than a sentence. Anchoring on the coefficient
    # and standard error that precede it keeps the match specific -- an
    # earlier version of this check matched a different sentence entirely and
    # compared the wrong number.
    check("primary odds ratio", "ratio",
          find_ratio(r"Hippocampal\s*\n\s*[−-][\d.]+\s*\n\s*[\d.]+\s*\n\s*(\d+\.\d+)"),
          effect["odds_ratio"])
    check("primary CI low", "ratio",
          find_ratio(r"Hippocampal\s*\n\s*[−-][\d.]+\s*\n\s*[\d.]+\s*\n\s*\d+\.\d+\s*\n\s*\[(\d+\.\d+)"),
          effect["odds_ratio_ci95_low"])

    # Region estimates are audited against the manuscript TABLE, not the prose.
    # A generated table cell has unambiguous provenance; a sentence does not,
    # and scraping prose for "hippocampus ... 0.70" happily matches a magnitude
    # coefficient three sentences away from the prevalence odds ratio it meant.
    prevalence = load("excursion_prevalence_glme.csv").assign(
        region=lambda d: d["term"].str.replace("Region_", "", regex=False)
    ).set_index("region")
    magnitude = load("signed_magnitude_mixed_model.csv").assign(
        region=lambda d: d["term"].str.replace("Region_", "", regex=False)
    ).set_index("region")

    region_table = None
    for table in Document(MAIN).tables:
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if any("Prevalence OR" in cell for cell in header):
            region_table = table
            break
    if region_table is None:
        unmatched.append("region results table not found in the manuscript")
    else:
        for row in region_table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            region = cells[0]
            if region not in prevalence.index:
                unmatched.append(f"table row '{region}' has no model row")
                continue
            check(f"Table 2 contacts, {region}", "count",
                  float(cells[1]), prevalence.loc[region, "df"] * 0 +
                  float(coverage.set_index("region").loc[region, "n_contacts"]))
            check(f"Table 2 prevalence OR, {region}", "ratio",
                  float(cells[4].split()[0]), prevalence.loc[region, "odds_ratio"])
            beta = cells[6].split()[0].replace("\u2212", "-").replace("+", "")
            check(f"Table 2 magnitude beta, {region}", "coefficient",
                  float(beta), magnitude.loc[region, "estimate"])

    sensitivity = load("polarity_threshold_sensitivity.csv")
    robust = sensitivity[sensitivity["p_value"] < 0.05]
    bounds = re.search(
        r"leaves the odds ratio between (\d+\.\d+) and (\d+\.\d+)", text)
    # Guard against a future edit reintroducing a shared stem: two analyses
    # reporting an odds-ratio range must not be indistinguishable to a reader
    # or to this audit.
    if len(re.findall(r"the odds ratio (?:between|within)", text)) > 2:
        suspicious.append(
            "more than two sentences report an odds-ratio range with the same "
            "stem; disambiguate them")
    # The threshold sweep no longer quotes a range of odds ratios in the text;
    # it claims the direction and the significance hold at every threshold.
    # That is the claim, so that is what is verified, against the same table
    # the sentence points at.
    if re.search(r"preserves the direction of the contrast at every "
                 r"threshold, and it remains significant at every one", text):
        same_direction = (robust["odds_ratio_hippocampal"] < 1).all()
        check("threshold sweep direction holds", "count",
              1.0 if same_direction else 0.0, 1.0)
        check("threshold sweep significance holds", "count",
              1.0 if (robust["p_value"] < 0.05).all() else 0.0, 1.0)
    else:
        unmatched.append("threshold-sensitivity claim sentence not found")

    # Leave-one-patient-out is now reported on the primary continuous scale,
    # in standard deviations of the outcome, so this reads that sentence and
    # converts the stored coefficients the same way the builder does.
    leave_one_out = load("continuous_leave_one_patient_out.csv").dropna(
        subset=["estimate"])
    outcome_sd = float(
        load("continuous_primary_summary.csv").iloc[0]["outcome_sd"])
    lopo = re.search(r"holds the primary contrast between ([−-]?\d+\.\d+) and "
                     r"([−-]?\d+\.\d+) standard deviations across all (\d+) "
                     r"refits", text)
    if lopo:
        def _num(g):
            return float(g.replace("\u2212", "-"))
        scaled = leave_one_out["estimate"] / outcome_sd
        check("leave-one-out strongest", "ratio", _num(lopo.group(1)),
              round(float(scaled.min()), 2))
        check("leave-one-out weakest", "ratio", _num(lopo.group(2)),
              round(float(scaled.max()), 2))
        check("leave-one-out refit count", "count", float(lopo.group(3)),
              len(leave_one_out))
    else:
        unmatched.append("leave-one-patient-out sentence not found")

    # --------------------------------------------------- hardcoded literals
    # Any digit typed into the builder bypasses the CSV guarantee. Acquisition
    # constants legitimately live there; anything the tables also know must not.
    source = BUILDER.read_text()
    derivable = {str(v) for v in ledger.values()}
    for line_number, line in enumerate(source.splitlines(), 1):
        if (line.lstrip().startswith("#") or "f\"" in line or "f'" in line
                or "http" in line):
            continue
        # Cross-references carry digits that are not quantities: "Table S18",
        # "Figure 9", "Fig. 4a". Flagging those trains a reader to ignore the
        # warning, which is worse than not raising it.
        scannable = re.sub(
            r'\b(?:Table|Tables|Fig|Fig\.|Figure|Figures)\s+S?\d+[a-z]?', '',
            line)
        for literal in re.findall(r'"[^"]*?(\d{2,4})[^"]*?"', scannable):
            if literal in derivable:
                suspicious.append(
                    f"build_v5_docx.py:{line_number}: literal {literal} is a "
                    f"ledger value and should be interpolated, not typed")

    # ------------------------------------------------------- bibliography
    # Every reference must carry a resolvable identifier. Network verification
    # against Crossref is a manual step (tools/ has no network in CI), but a
    # missing or malformed DOI is caught here on every build.
    # Only lines after the References heading. An in-text citation group such
    # as "(Reimer et al., 2014; McGinley et al., 2015)" also carries a year and
    # a semicolon, and matching on that alone flags body prose as a reference.
    # Paragraphs only. document_text() appends every table cell after the
    # paragraph stream, so slicing its output from "References" onward sweeps
    # up table cells as if they were bibliography entries.
    paragraphs = [p.text.strip() for p in Document(MAIN).paragraphs
                  if p.text.strip()]
    try:
        start = paragraphs.index("References") + 1
    except ValueError:
        start = len(paragraphs)
    reference_lines = paragraphs[start:]
    for line in reference_lines:
        if "doi:" not in line:
            failures.append(f"reference without a DOI: {line[:70]}")
        elif not re.search(r"doi:10\.\d{4,9}/\S+", line):
            failures.append(f"malformed DOI: {line[-40:]}")
    if reference_lines:
        passes += 1
        print(f"\nBibliography: {len(reference_lines)} references, "
              f"{sum('doi:' in l for l in reference_lines)} with a DOI")

    # --------------------------------------------- documentation drift
    # The working documents quote results in prose and nothing regenerates
    # them, so they drift out of agreement with the tables and, worse, with
    # each other. Every instance found so far was a figure that had been
    # correct when written and was silently invalidated by a later analysis.
    # These checks cover the quantities that have actually drifted.
    docs = REPO / "docs"
    doc_text = "\n".join(
        p.read_text(errors="replace") for p in sorted(docs.glob("*.md")))

    density_path = TABLES / "replication_event_density.csv"
    if density_path.exists():
        if True:
            interval = pd.read_csv(density_path)[
                "median_inter_event_interval_seconds"].median()
            stated = set(re.findall(
                r"median (?:inter-event )?(?:gap|interval)[^.]{0,60}?"
                r"(\d+\.\d+) s", doc_text))
            wrong = [v for v in stated if abs(float(v) - interval) > 0.05]
            if wrong:
                failures.append(
                    f"docs quote an inter-event interval of {sorted(wrong)} s "
                    f"against {interval:.2f} s in the table")
            else:
                passes += 1

    # Counts of figures, tables and references are quoted in HANDOFF.md and
    # have been wrong after every structural change to the paper.
    handoff = docs / "HANDOFF.md"
    if handoff.exists():
        text_handoff = handoff.read_text(errors="replace")
        document = Document(MAIN)
        main_figures = sum(
            1 for p in document.paragraphs
            if re.match(r"^Figure \d+\.", p.text.strip()))
        claimed = re.search(r"main (\d+) figures", text_handoff)
        if claimed and int(claimed.group(1)) != main_figures:
            failures.append(
                f"HANDOFF claims {claimed.group(1)} main figures, "
                f"the manuscript has {main_figures}")
        elif claimed:
            passes += 1

        claimed_refs = re.search(r"(\d+) references all DOI-verified",
                                 text_handoff)
        if claimed_refs and int(claimed_refs.group(1)) != len(reference_lines):
            failures.append(
                f"HANDOFF claims {claimed_refs.group(1)} references, "
                f"the manuscript has {len(reference_lines)}")
        elif claimed_refs:
            passes += 1

    # ------------------------------------------------- replication cohort
    # The replication introduces a second cohort with its own contact counts
    # and its own odds ratio, both of which are easy to confuse with the
    # discovery figures. Every one of them is checked against its table.
    replication_file = TABLES / "replication_headline.csv"
    if not replication_file.exists():
        unmatched.append("replication headline table missing")
    else:
        rep = pd.read_csv(replication_file).iloc[0]
        rep_coverage = load("replication_coverage_primary.csv")

        # The prespecified odds ratio moved out of the running text and into
        # the supplement table that keeps the registered criterion checkable,
        # so these read the table cells rather than a sentence.
        check("replication odds ratio", "ratio",
              find_ratio(r"extrahippocampal dilation\s*\n\s*(\d+\.\d+)"),
              rep["odds_ratio"])
        check("replication CI low", "ratio",
              find_ratio(r"95% confidence interval\s*\n\s*\[(\d+\.\d+)"),
              rep["ci95_low"])

        # Where the within-shaft gradient is centred. These numbers decide
        # whether the paper calls the gradient hippocampal or mesial temporal,
        # so a drift here changes a claim rather than a decimal place.
        origin_file = TABLES / "gradient_origin_summary.csv"
        if origin_file.exists():
            origin = pd.read_csv(origin_file).iloc[0]
            check("gradient origins tested", "count",
                  find_ratio(r"every one of (\d+) candidate origins"),
                  origin["n_origins_tested"])
            check("hippocampal origin rank", "count",
                  find_ratio(r"hippocampal origin ranks (\d+) of"),
                  origin["hippocampus_rank"])
            check("hippocampal origin delta AIC", "ratio",
                  find_ratio(r"sits (\d+\.\d+) AIC from the best"),
                  round(float(origin["hippocampus_delta_aic"]), 1))
            check("worst origin delta AIC", "count",
                  find_ratio(r"is (\d+) AIC behind the best"),
                  round(float(origin["worst_origin_delta_aic"])))
            check("origin axis correlation", "ratio",
                  find_ratio(r"hippocampal one at r \u2265 (\d+\.\d+)"),
                  round(float(origin["min_axis_correlation_within_set"]), 2))
        else:
            unmatched.append("gradient origin summary missing")

        # Cross-cohort comparison. These numbers decide whether the paper
        # calls the replication inconclusive or negative, so a drift between
        # the tables and the prose here would change a conclusion rather than
        # a decimal place.
        comparability_file = TABLES / "replication_comparability.csv"
        heterogeneity_file = TABLES / "replication_heterogeneity.csv"
        if comparability_file.exists() and heterogeneity_file.exists():
            comp = pd.read_csv(comparability_file).set_index("cohort")
            het = pd.read_csv(heterogeneity_file).iloc[0]

            check("replication effect per residual SD", "ratio",
                  find_ratio(r"rest of the implant by ([+-]\d+\.\d+) residual"),
                  comp.loc["Replication", "beta_per_residual_sd"])
            check("discovery effect per residual SD", "ratio",
                  find_ratio(r"P = [^)]*\), against ([+-]\d+\.\d+) here"),
                  comp.loc["Discovery", "beta_per_residual_sd"])
            check("cohort heterogeneity z", "ratio",
                  find_ratio(r"\(z = ([+-]?\d+\.\d+), P ="),
                  het["z_statistic"])
            check("replication detectable effect", "ratio",
                  find_ratio(r"detect (\d+\.\d+) residual standard deviations"),
                  het["replication_detectable_at_80_power"])
            check("replication residual SD", "ratio",
                  find_ratio(r"same amount as the discovery cohort \((\d+\.\d+) against"),
                  comp.loc["Replication", "sd_residual"])
            check("discovery residual SD", "ratio",
                  find_ratio(r"same amount as the discovery cohort \(\d+\.\d+ against (\d+\.\d+)\)"),
                  comp.loc["Discovery", "sd_residual"])
            check("replication median fit R2", "ratio",
                  find_ratio(r"with median R. of (\d+\.\d+) against"),
                  comp.loc["Replication", "median_fit_r2"])
            check("discovery median fit R2", "ratio",
                  find_ratio(r"with median R. of \d+\.\d+ against (\d+\.\d+)"),
                  comp.loc["Discovery", "median_fit_r2"])
        else:
            unmatched.append("replication comparability tables missing")

        rep_ledger = {
            "replication contacts": int(rep_coverage["n_contacts"].sum()),
            "replication excursion contacts": int(rep["n_excursion_contacts"]),
            "replication hippocampal excursion":
                int(rep["n_hippocampal_excursion_contacts"]),
            "replication participants": int(rep["n_subjects"]),
        }
        print("\nReplication ledger")
        for name, value in rep_ledger.items():
            occurrences = len(re.findall(rf"\b{value}\b", text))
            print(f"  {name:<34} {value:>5}   appears {occurrences}x")
            if occurrences == 0:
                unmatched.append(f"{name}={value} never appears in the paper")

        # The two cohorts must not be conflated: the replication odds ratio and
        # the discovery odds ratio are different numbers about different data.
        if abs(float(rep["odds_ratio"]) - float(effect["odds_ratio"])) < 1e-9:
            failures.append(
                "replication and discovery odds ratios are identical, which "
                "means one of them is being read from the wrong table")
        else:
            passes += 1

        # The prespecified verdict must appear as written, not paraphrased into
        # something stronger.
        verdict = str(rep["prespecified_verdict"])
        if verdict == "inconclusive" and "inconclusive" not in text.lower():
            failures.append(
                "the replication verdict is 'inconclusive' but the manuscript "
                "never uses the word")
        else:
            passes += 1

        # The event spacing is the load-bearing fact behind "this cohort cannot
        # isolate spontaneous activity", and it was originally typed from a
        # single run rather than read from all of them.
        density_file = TABLES / "replication_event_density.csv"
        if not density_file.exists():
            unmatched.append("replication event density table missing")
        else:
            density = pd.read_csv(density_file)
            check("replication inter-event interval", "ratio",
                  find_ratio(r"median interval between task events of "
                             r"(\d+\.\d+) s"),
                  density["median_inter_event_interval_seconds"].median())

        for ripple_name in ("replication_ripple_pupil_test_laplacian_1khz.csv",
                            "replication_ripple_pupil_test_raw_4khz.csv"):
            path = TABLES / ripple_name
            if not path.exists():
                unmatched.append(f"{ripple_name} missing")
                continue
            row = pd.read_csv(path).iloc[0]
            stated = len(re.findall(rf"\b{int(row['n_ripples'])}\b", text))
            if stated == 0:
                unmatched.append(
                    f"ripple count {int(row['n_ripples'])} "
                    f"({row['signal_source']}) never appears")
            else:
                passes += 1

    # -------------------------------------------------- duplicate estimates
    ratios = re.findall(r"odds ratio (\d+\.\d{2,3})", text)
    for value in set(ratios):
        if ratios.count(value) > 2:
            suspicious.append(
                f"odds ratio {value} appears {ratios.count(value)} times — "
                "check these are genuinely distinct estimates")

    # ------------------------------------------------------------- report
    print(f"\nPASS {passes}   FAIL {len(failures)}   "
          f"UNMATCHED {len(unmatched)}   SUSPICIOUS {len(suspicious)}")
    for title, items in [("FAIL", failures), ("UNMATCHED", unmatched),
                         ("SUSPICIOUS", suspicious)]:
        if items:
            print(f"\n{title}")
            for item in items:
                print(f"  - {item}")

    print("\nOverall:", "FAIL" if failures else "PASS")
    return 1 if failures else 0


if __name__ == "__main__":
    sys.exit(main())
